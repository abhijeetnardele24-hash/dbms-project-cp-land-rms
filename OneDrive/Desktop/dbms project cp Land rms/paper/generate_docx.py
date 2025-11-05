# Generates an IEEE-style DOCX paper for the LRMS project
# Usage: python generate_docx.py

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.abspath(__file__))
REFS_PATH = os.path.join(ROOT, 'refs.bib')
OUTPUT_PATH = os.path.join(ROOT, 'lrms_ieee.docx')


def set_normal_style(document):
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    # Ensure Times New Roman in underlying rPr
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def add_title_and_author(document):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('An Enterprise-Grade Land Registry Management System Using Flask and MySQL with Advanced Database Features')
    run.bold = True
    run.font.size = Pt(16)

    author = document.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.add_run('Abhijeet Nardele').bold = True

    aff = document.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff.add_run('Department of Computer Engineering, [Your Institution Name]\n')
    aff.add_run('Email: your.email@example.com')


def add_abstract_and_keywords(document):
    p = document.add_paragraph()
    p.add_run('Abstract—').bold = True
    p.add_run(
        'Digitizing land records is critical for governance, transparency, and citizen services. '
        'This paper presents the design and implementation of a Land Registry Management System (LRMS) built with '
        'Flask (Python) and MySQL. The system supports role-based workflows across property registration, ownership, '
        'mutations, tax assessment and payments, document management, notifications, and audit logging. '
        'A key contribution is the use of advanced MySQL features—stored procedures, triggers, views, and strategic indexing—'
        'to enforce business rules and improve performance. We integrate interactive mapping with Leaflet and analytics with Chart.js. '
        'We describe the architecture, schema, security controls, implementation details, and validation on a realistic dataset.'
    )

    k = document.add_paragraph()
    k.add_run('Index Terms—').bold = True
    k.add_run('Land registry, DBMS, Flask, MySQL 8.0, Stored procedures, Triggers, Views, GIS, Leaflet, RBAC, E-governance')


def add_heading(document, text):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def add_sections(document):
    # Introduction
    add_heading(document, 'I. INTRODUCTION')
    document.add_paragraph(
        'We present LRMS, a web-based system implementing end-to-end workflows: property registration with ULPIN generation, '
        'ownership and joint-ownership, mutation processing, tax assessment and payments, document management, '
        'notifications, and comprehensive audit trails.'
    )
    document.add_paragraph('Contributions:')
    document.add_paragraph('• Role-based LRMS in Flask/SQLAlchemy', style=None)
    document.add_paragraph('• Advanced MySQL layer (procedures, triggers, views, indexes)', style=None)
    document.add_paragraph('• GIS with Leaflet; analytics with Chart.js', style=None)
    document.add_paragraph('• Security hardening (hashing, CSRF, RBAC, uploads, auditing)', style=None)

    # Background
    add_heading(document, 'II. BACKGROUND AND RELATED WORK')
    document.add_paragraph(
        'Government initiatives adopt ULPIN and digitized registries to enhance transparency. '
        'Our approach emphasizes moving business logic into the database via stored programs and triggers for '
        'consistency and auditability, with a clean application layer.'
    )

    # System Overview
    add_heading(document, 'III. SYSTEM OVERVIEW AND ARCHITECTURE')
    document.add_paragraph(
        'Three-tier architecture: Flask web application (Blueprints for roles and APIs), MySQL 8.0 database with advanced '
        'features, and browser client with Bootstrap, Chart.js, and Leaflet.'
    )
    document.add_paragraph('Figure 1 (placeholder): LRMS architecture diagram.')

    # Database Design
    add_heading(document, 'IV. DATABASE DESIGN')
    document.add_paragraph(
        'Core entities: users, properties, owners, ownerships (many-to-many), mutations, payments, documents, '
        'notifications, audit_logs, tax_assessments. Constraints and indexes ensure integrity and performance.'
    )
    document.add_paragraph('Figure 2 (placeholder): ER diagram for LRMS.')

    # Advanced MySQL Features
    add_heading(document, 'V. ADVANCED MYSQL FEATURES')
    document.add_paragraph('Stored Procedures: calculate_property_tax, get_property_report, get_ownership_chain, get_dashboard_stats.')
    document.add_paragraph('Triggers: after_property_insert, before_property_update, after_payment_insert, after_property_status_update.')
    document.add_paragraph('Views and Indexing: dashboard, revenue, geographic summaries; composite and full-text indexes.')

    # Methodology
    add_heading(document, 'VI. METHODOLOGY AND WORKFLOW')
    document.add_paragraph(
        'Workflow: (1) citizen registers property with geolocation; (2) registrar reviews and approves; '
        '(3) citizen submits mutation; (4) officer verifies and decides; (5) tax assessed and payment recorded; '
        '(6) notifications and audit logs generated.'
    )

    # Application Layer
    add_heading(document, 'VII. APPLICATION LAYER')
    document.add_paragraph(
        'Flask Blueprints (admin, registrar, officer, citizen, auth, api) with SQLAlchemy ORM, Flask-Login, WTForms. '
        'UI uses Bootstrap 5, Chart.js, and Leaflet for interactive maps; documents validated; reports exportable to PDF/Excel.'
    )

    # Security
    add_heading(document, 'VIII. SECURITY')
    document.add_paragraph(
        'Controls include password hashing, CSRF protection, session hardening, role-based decorators, strict file-type checks, '
        'and comprehensive audit logging. ORM mitigates injection; sensitive operations are logged with actor and timestamp.'
    )

    # Evaluation
    add_heading(document, 'IX. EVALUATION')
    document.add_paragraph('Functional validation across roles and database-level enforcement were verified on seeded data.')
    # Table of DB objects
    table = document.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Object'
    hdr_cells[1].text = 'Count/Examples'
    data = [
        ('Stored procedures', '4 (tax, report, ownership chain, dashboard stats)'),
        ('Triggers', '4 (property insert/update, payment insert, status update)'),
        ('Views', '7 (dashboard, revenue, geographic, owners)'),
        ('Indexes', '12+ (status, ULPIN, location, payments)'),
        ('Tables', '10+ core (users, properties, owners, ...)'),
    ]
    for i, (k, v) in enumerate(data, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = k
        row_cells[1].text = v

    # Discussion
    add_heading(document, 'X. DISCUSSION AND LIMITATIONS')
    document.add_paragraph(
        'Production deployment benefits from partitioning high-volume tables, read replicas, Redis caching, and '
        'hardened secrets management.'
    )

    # Future Work
    add_heading(document, 'XI. FUTURE WORK')
    document.add_paragraph(
        'Integrate real payment gateway, OCR for documents, digital signatures, scheduled events, geographic sharding, '
        'advanced GIS layers.'
    )

    # Conclusion
    add_heading(document, 'XII. CONCLUSION')
    document.add_paragraph(
        'LRMS demonstrates a secure, auditable, and performant e-governance system with business logic in MySQL '
        'and a clean Flask service layer.'
    )

    # Ethics & Acknowledgment
    add_heading(document, 'Ethical Considerations and Data Availability')
    document.add_paragraph('No real PII was used; all experiments were done on seeded/test data. Materials available with the source.')
    add_heading(document, 'ACKNOWLEDGMENT')
    document.add_paragraph('Completed as part of a DBMS course project; thanks to mentors and peers for feedback.')


def parse_refs():
    if not os.path.exists(REFS_PATH):
        return []
    entries = []
    with open(REFS_PATH, 'r', encoding='utf-8') as f:
        block = []
        for line in f:
            if line.strip().startswith('@'):
                if block:
                    entries.append('\n'.join(block))
                    block = []
            block.append(line.rstrip())
        if block:
            entries.append('\n'.join(block))
    formatted = []
    for e in entries:
        def get_field(name):
            import re
            m = re.search(r'\b' + name + r'\s*=\s*\{([^}]*)\}', e)
            return m.group(1) if m else ''
        title = get_field('title') or 'Untitled'
        author = get_field('author') or get_field('howpublished')
        url = get_field('url')
        year = get_field('year')
        parts = [title]
        if author:
            parts.append(author)
        if year:
            parts.append(year)
        if url:
            parts.append(url)
        formatted.append('; '.join(parts))
    return formatted


def add_references(document):
    add_heading(document, 'REFERENCES')
    refs = parse_refs()
    if not refs:
        document.add_paragraph('[1] Flask Documentation, https://flask.palletsprojects.com/')
        return
    for i, ref in enumerate(refs, start=1):
        document.add_paragraph(f'[{i}] {ref}')


def main():
    document = Document()
    set_normal_style(document)
    add_title_and_author(document)
    add_abstract_and_keywords(document)
    add_sections(document)
    add_references(document)
    document.save(OUTPUT_PATH)
    print(f'Wrote {OUTPUT_PATH}')


if __name__ == '__main__':
    main()
